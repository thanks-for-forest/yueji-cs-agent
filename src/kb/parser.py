"""知识库文档解析器：支持 .md / .docx / .pdf / .txt / .csv / .xlsx / .html → 纯文本 + 分块。

借鉴 Langchain-Chatchat / RAGFlow 的 KB 解析模式，轻量自研。
"""
from __future__ import annotations

import re
from pathlib import Path

SUPPORTED_EXT = {".md", ".markdown", ".docx", ".pdf", ".txt", ".csv", ".xlsx", ".html", ".htm"}


def parse_document(filename: str, data: bytes) -> str:
    """按扩展名解析文档为纯文本。"""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXT:
        raise ValueError(f"不支持的格式：{ext}（支持 {', '.join(sorted(SUPPORTED_EXT))}）")
    if ext in (".md", ".markdown", ".txt"):
        return data.decode("utf-8", errors="replace")
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".csv":
        return _parse_csv(data)
    if ext == ".xlsx":
        return _parse_xlsx(data)
    if ext in (".html", ".htm"):
        return _parse_html(data)
    raise ValueError(f"不支持的格式：{ext}")


def _parse_docx(data: bytes) -> str:
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def _parse_csv(data: bytes) -> str:
    """CSV → 每行转 '列1 | 列2 | …' 文本。"""
    import csv
    import io

    rows = list(csv.reader(io.StringIO(data.decode("utf-8-sig", errors="replace"))))
    if not rows:
        return ""
    header = rows[0]
    parts = [" | ".join(header)]
    for row in rows[1:]:
        # 跳过全空行
        if any(cell.strip() for cell in row):
            parts.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    """xlsx → 每个 sheet 转行文本（借用 openpyxl）。"""
    import io

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        parts.append(f"【工作表 {ws.title}】")
        for row in rows:
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _parse_html(data: bytes) -> str:
    """HTML → 剥标签与脚本/样式，保留正文文本。"""
    from lxml import html as lx

    root = lx.fromstring(data.decode("utf-8", errors="replace"))
    # 移除脚本/样式
    for tag in root.iter("script", "style", "noscript"):
        tag.drop_tree()
    text = root.text_content()
    # 压缩空白
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 60) -> list[str]:
    """按字符窗口分块（保留语义边界：优先在段落/句号处切断）。"""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        if end < n:
            # 在窗口内找最近的段落/句号边界
            window = text[start:end]
            cut = -1
            for sep in ("\n\n", "\n", "。", "！", "？", "；"):
                idx = window.rfind(sep)
                if idx > chunk_size * 0.4:
                    cut = idx
                    break
            if cut > 0:
                end = start + cut + len(sep)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break  # 已到文末，收尾
        start = end - overlap
    return chunks
